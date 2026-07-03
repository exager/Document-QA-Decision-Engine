import logging
from io import BytesIO
from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.core.documents.elements import (
    DocumentElement,
    ElementType,
    ExtractedDocument,
)
from app.core.loaders.base import DocumentLoader

logger = logging.getLogger(__name__)

# python-docx style names → element types.
_STYLE_MAP = {
    "Title": ElementType.TITLE,
    "Heading 1": ElementType.HEADING,
    "Heading 2": ElementType.HEADING,
    "Heading 3": ElementType.HEADING,
    "Heading 4": ElementType.HEADING,
    "Heading 5": ElementType.HEADING,
    "Heading 6": ElementType.HEADING,
    "List Paragraph": ElementType.LIST_ITEM,
    "List Bullet": ElementType.LIST_ITEM,
    "List Number": ElementType.LIST_ITEM,
    "Caption": ElementType.CAPTION,
}

class DocxLoader(DocumentLoader):
    def load(self, file_bytes: bytes) -> ExtractedDocument:
        try:
            doc = DocxDocument(BytesIO(file_bytes))
        except Exception as exc:
            logger.error("docx_open_failed", extra={"error": str(exc)})
            return ExtractedDocument(
                elements=[],
                metadata={"loader": "docx", "warnings": [f"open_failed: {exc}"]},
            )

        elements: list[DocumentElement] = []
        warnings: list[str] = []
        heading_path: list[str] = []

        # Walk the document body in reading order: paragraphs AND tables interleaved.
        # python-docx doesn't expose this directly; we walk the XML element children.
        for block in _iter_block_items(doc):
            if isinstance(block, Paragraph):
                el = _paragraph_to_element(block, heading_path)
                if el is None:
                    continue
                style_name = el.metadata.get("style", "") 
                if el.type in (ElementType.TITLE, ElementType.HEADING):
                    heading_path = _push_heading(heading_path, el.text, style_name)
                    el = DocumentElement(
                        type=el.type,
                        text=el.text,
                        heading_path=tuple(h for h in heading_path if h),
                        metadata=el.metadata,
                    )
                elements.append(el)

            elif isinstance(block, Table):
                try:
                    md = _table_to_markdown(block)
                    if md:
                        elements.append(
                            DocumentElement(
                                type=ElementType.TABLE,
                                text=md,
                                heading_path=tuple(heading_path),
                                metadata={
                                    "rows": len(block.rows),
                                    "cols": len(block.columns),
                                },
                            )
                        )
                except Exception as exc:
                    warnings.append(f"table_extract_failed: {exc}")

        return ExtractedDocument(
            elements=elements,
            metadata={
                "loader": "docx",
                "element_count": len(elements),
                "warnings": warnings,
            },
        )


def _iter_block_items(doc):
    from docx.oxml.ns import qn

    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def _paragraph_to_element(p: Paragraph, heading_path: list[str]) -> DocumentElement | None:
    from docx.oxml.ns import qn
    
    text = p.text.strip()
    if not text:
        return None

    style_name = (p.style.name if p.style else "") or ""
    etype = _STYLE_MAP.get(style_name, ElementType.PARAGRAPH)

    if p._element.find(qn("w:numPr")) is not None:
        etype = ElementType.LIST_ITEM
    return DocumentElement(
        type=etype,
        text=text,
        heading_path=tuple(heading_path),
        metadata={"style": style_name},
    )

def _table_to_markdown(table: Table) -> str:
    rows: list[list[str]] = []
    rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
    if not rows or not rows[0]:
        return ""
    header, *body = rows
    lines = ["| " + " | ".join(header) + " |",
             "| " + " | ".join(["---"] * len(header)) + " |"]
    for r in body:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


def _push_heading(stack: list[str], new_heading: str, style: str) -> list[str]:
    """
    Use the Heading N number to place the new heading at level N.
    Everything deeper than N gets popped.
    """
    level = 1
    if style.startswith("Heading "):
        try:
            level = int(style.split()[1])
        except (IndexError, ValueError):
            level = 1
    elif style == "Title":
        level = 1 
    stack = stack[: level - 1]
    if len(stack) < level - 1:
        stack = stack + [""] * (level - 1 - len(stack))
    stack.append(new_heading)
    return stack